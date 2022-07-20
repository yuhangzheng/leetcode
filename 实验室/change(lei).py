from docx import Document
from docx.shared import Inches, Pt
from win32com.client import Dispatch, constants, gencache
import os
import re
import copy
import time

names = input("请输入姓名(空格隔开)(回车确定)：")
real_path = os.getcwd()
files = os.listdir(real_path)
name_list = names.split(" ")


def transfer(path, file, name):

    word_path = path + '\\' + file
    pdf_path = path + "\\" + name + "欢迎信.pdf"
    document = Document(word_path)
    print(word_path)
    for para in document.paragraphs:
        if re.match(r"亲爱的", para.text):
            # 深度复制段落中的内容
            list_runs = copy.deepcopy(para.runs)
            # 清空段落中的内容
            para.clear()
            # 遍历已复制的段落中的内容
            for run in list_runs:
                if run.text == "X":
                    run_text = para.add_run(name[0])
                    run_text.font.name = u'宋体'
                    run_text.font.size = Pt(10.5)
                    run_text.font.underline = True
                elif run.text == "XX":
                    run_text = para.add_run(name[1:])
                    run_text.font.name = u'宋体'
                    run_text.font.size = Pt(10.5)
                    run_text.font.underline = True
                else:
                    for t in run.text:
                        if t == " ":
                            run_text = para.add_run(t)
                            run_text.font.underline = True
                        else:
                            run_text = para.add_run(t)
    document.save(word_path)

    # word 转 pdf
    wd = gencache.EnsureDispatch('Word.Application')
    doc = wd.Documents.Open(word_path)
    doc.ExportAsFixedFormat(pdf_path,
                            constants.wdExportFormatPDF,
                            Item=constants.wdExportDocumentWithMarkup,
                            CreateBookmarks=constants.wdExportCreateHeadingBookmarks)
    print("{}转换完成".format(file))
    wd.Quit(constants.wdDoNotSaveChanges)


    time.sleep(5)
    # 还原模块
    for para in document.paragraphs:
        if re.match(r"亲爱的", para.text):
            para.clear()

            for run in list_runs:
                if run.text == "X":
                    run_text = para.add_run("X")
                    run_text.font.underline = True
                elif run.text == "XX":
                    run_text = para.add_run("XX")
                    run_text.font.underline = True
                else:
                    for t in run.text:
                        if t == " ":
                            run_text = para.add_run(t)
                            run_text.font.underline = True
                        else:
                            run_text = para.add_run(t)
    document.save(word_path)


for i, list_name in enumerate(name_list):
    for real_file in files:
        if real_file.split('.')[-1] in ['docx', 'doc']:
            try:
                transfer(real_path, real_file, list_name)
            except Exception as e:
                print(" ")

